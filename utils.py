import csv
import io
import re

from docx import Document

# Ordered most-specific first so "cover letter" matches before bare "letter"
_DOC_PATTERNS: list[tuple[list[str], str]] = [
    (["cover letter"],                        "cover_letter.docx"),
    (["resignation letter", "resignation"],   "resignation_letter.docx"),
    (["thank you letter", "thank-you"],       "thank_you_letter.docx"),
    (["press release"],                       "press_release.docx"),
    (["resume", " cv "],                      "resume.docx"),
    (["letter"],                              "letter.docx"),
    (["email", "e-mail"],                     "email.docx"),
    (["proposal"],                            "proposal.docx"),
    (["biography", " bio "],                  "bio.docx"),
    (["speech"],                              "speech.docx"),
    (["essay"],                               "essay.docx"),
    (["report"],                              "report.docx"),
    (["memo"],                                "memo.docx"),
    (["announcement"],                        "announcement.docx"),
]


def detect_document_type(prompt: str) -> tuple[bool, str]:
    """
    Return (is_document, filename) based on keywords in the user prompt.
    Used to decide whether to show a .txt download button.
    """
    lower = f" {prompt.lower()} "   # pad so word-boundary checks work
    for keywords, filename in _DOC_PATTERNS:
        if any(kw in lower for kw in keywords):
            return True, filename
    return False, ""


def extract_code_block(text: str) -> tuple[str, str]:
    """Extract the first fenced code block. Returns (code, language)."""
    match = re.search(r"```(\w*)\n(.*?)```", text, re.DOTALL)
    if match:
        lang = match.group(1).strip() or "text"
        code = match.group(2).rstrip()
        return code, lang
    return "", ""


def detect_language(fence_hint: str) -> str:
    """Normalize a code-fence language hint to a Streamlit-recognized token."""
    _map = {
        "py": "python",
        "js": "javascript",
        "ts": "typescript",
        "sh": "bash",
        "shell": "bash",
        "zsh": "bash",
        "rb": "ruby",
        "rs": "rust",
        "": "text",
    }
    hint = fence_hint.lower().strip()
    return _map.get(hint, hint or "text")


def text_to_docx(text: str) -> bytes:
    """Convert a plain-text / markdown-lite string to a .docx file (bytes)."""
    doc = Document()
    for line in text.splitlines():
        line = line.strip()
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line:
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def extract_markdown_table_to_csv(text: str) -> str:
    """
    Find the first markdown table in text and return it as a CSV string.
    Returns an empty string if no table is found.
    """
    table_lines = []
    in_table = False
    for line in text.splitlines():
        if "|" in line:
            in_table = True
            table_lines.append(line)
        elif in_table:
            break

    if not table_lines:
        return ""

    rows = []
    for line in table_lines:
        if re.match(r"^\|[\s\-:|]+\|$", line.strip()):
            continue  # skip separator row
        cells = [c.strip() for c in line.strip().split("|")]
        cells = [c for c in cells if c]
        if cells:
            rows.append(cells)

    if not rows:
        return ""

    out = io.StringIO()
    csv.writer(out).writerows(rows)
    return out.getvalue()


def pre_code_text(text: str) -> str:
    """Return the text that appears before the first code fence."""
    idx = text.find("```")
    if idx == -1:
        return text.strip()
    return text[:idx].strip()
